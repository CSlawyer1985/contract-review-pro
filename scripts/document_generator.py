"""
文档生成模块（V4.0）

终稿三件套真正的 .docx 生成（旧版 .md 输出已移除——违反"禁止 .md 终稿"硬约束）：
- 法律意见书 / 法律分析：python-docx 从零创建（新文档合法场景），版式规范内置
  深蓝标题 + 仿宋正文 + 浅底元信息卡 + 棕色标签 + 页脚页码 + 紧凑正式件参数
- 批注版合同：走 docx_generator 双引擎（docx skill 增强 / ooxml_lite 内置兜底）

版式常量集中在模块顶部，可按律所模板统一调整。
"""

from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# 版式规范（C2）
# ---------------------------------------------------------------------------
CLR_HEADING = RGBColor(0x1F, 0x38, 0x64)   # 深蓝标题
CLR_LABEL = RGBColor(0x97, 0x48, 0x06)     # 棕色标签高亮
CLR_HIGH = RGBColor(0xC0, 0x00, 0x00)      # 高风险红
CLR_MID = RGBColor(0xC5, 0x50, 0x0B)       # 中风险橙
CLR_LOW = RGBColor(0x2E, 0x74, 0xB5)       # 低风险蓝
SHADE_CARD = "DCE6F1"                       # 元信息卡浅蓝底
SHADE_HEAD = "F2F2F2"                       # 表头浅灰底

FONT_HEADING_EAST = "黑体"
FONT_BODY_EAST = "仿宋"
FONT_WEST = "Times New Roman"

LEVEL_COLOR = {"致命风险": CLR_HIGH, "重要风险": CLR_MID,
               "一般风险": CLR_LOW, "轻微瑕疵": CLR_LOW}
LEVEL_TEXT = {"致命风险": "高", "重要风险": "中", "一般风险": "低", "轻微瑕疵": "低"}


class DocumentGenerator:
    """文档生成器（V4.0 docx 原生输出）"""

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # =======================================================================
    # 版式基础设施
    # =======================================================================

    def _base_doc(self) -> Document:
        """创建应用版式规范的空白文档：紧凑页边距 + 仿宋正文 + 深蓝标题 + 页脚页码"""
        doc = Document()

        # 紧凑正式件参数：压缩页边距/行距/段距，避免无效留白
        for section in doc.sections:
            section.top_margin = Cm(2.2)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            self._add_page_number(section)

        normal = doc.styles["Normal"]
        normal.font.name = FONT_WEST
        normal.font.size = Pt(11)
        normal.element.get_or_add_rPr()
        rfonts = normal.element.rPr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), FONT_BODY_EAST)
        pf = normal.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(20)
        pf.space_after = Pt(4)

        for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
            st = doc.styles[name]
            st.font.name = FONT_WEST
            st.font.size = Pt(size)
            st.font.color.rgb = CLR_HEADING
            st.font.bold = True
            st.element.get_or_add_rPr()
            st.element.rPr.get_or_add_rFonts().set(qn("w:eastAsia"), FONT_HEADING_EAST)
            st.paragraph_format.space_before = Pt(10)
            st.paragraph_format.space_after = Pt(6)

        return doc

    @staticmethod
    def _add_page_number(section):
        """页脚居中页码（PAGE 域）"""
        para = section.footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), " PAGE ")
        run._r.append(fld)

    @staticmethod
    def _shade(cell, hex_color: str):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shd)

    def _meta_card(self, doc: Document, rows: List[tuple]):
        """浅底元信息卡：单列表格，左键右值"""
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (key, value) in enumerate(rows):
            k_cell, v_cell = table.rows[i].cells
            k_cell.width = Cm(3.2)
            v_cell.width = Cm(12.8)
            self._shade(k_cell, SHADE_CARD)
            k_run = k_cell.paragraphs[0].add_run(key)
            k_run.font.bold = True
            k_run.font.color.rgb = CLR_HEADING
            v_cell.paragraphs[0].add_run(str(value))
        return table

    @staticmethod
    def _label_run(paragraph, text: str, color: RGBColor, bold: bool = True):
        run = paragraph.add_run(text)
        run.font.color.rgb = color
        run.font.bold = bold
        return run

    # =======================================================================
    # 1. 法律意见书（客户版 · 五模块 + 客户速览页）
    # =======================================================================

    def generate_legal_opinion_docx(self, contract_name: str, analysis_result: Dict,
                                     risk_report: Dict, user_context: Dict,
                                     author: str = "陈石律师【海泰所】") -> str:
        """生成法律意见书 .docx（客户速览页 + 五模块正式版）"""
        filepath = self.output_dir / f"{contract_name}-法律意见书.docx"

        risks_by_level = risk_report.get("risks_by_level", {})
        summary = risk_report.get("summary", {})
        radar_data = risk_report.get("radar_data", {})
        label_dist = risk_report.get("label_distribution", {})
        scoring = risk_report.get("scoring", {})
        today = datetime.now().strftime("%Y年%m月%d日")

        doc = self._base_doc()

        # ------------------------------------------------------------------
        # 客户速览页（首页 1 页内：结论段 + Top 风险 + 谈判要点）
        # ------------------------------------------------------------------
        doc.add_heading(f"{contract_name} 法律审核意见书", level=0).alignment = \
            WD_ALIGN_PARAGRAPH.CENTER

        self._meta_card(doc, [
            ("合同名称", contract_name),
            ("合同类型", analysis_result.get("identified_type", "未识别")),
            ("委托方立场", user_context.get("party", "未指定")),
            ("审核日期 / 审核人", f"{today} / {author}"),
        ])

        doc.add_heading("审核结论速览", level=1)
        brief = user_context.get("client_brief")
        if brief:
            doc.add_paragraph(brief)
        else:
            p = doc.add_paragraph()
            self._label_run(p, "【结论段占位】", CLR_LABEL)
            doc.add_paragraph(
                '以 3-5 句话概括：合同整体可签性、最重要的 1-2 个风险、最优先的谈判动作。'
                '遵循“律师分析利弊，客户做决定”原则，不写“建议签署/不建议签署”的刚性结论。')

        # Top 风险（≤5 条，零法条术语）
        doc.add_heading("最需要关注的风险（Top 5）", level=2)
        top_risks = (risks_by_level.get("致命风险", []) +
                     risks_by_level.get("重要风险", []))[:5]
        if top_risks:
            for i, risk in enumerate(top_risks, 1):
                p = doc.add_paragraph()
                self._label_run(p, f"{i}. ", CLR_LABEL)
                p.add_run(risk.get("description", "")[:80])
        else:
            doc.add_paragraph("未发现致命/重要级别风险。")

        # 谈判要点
        doc.add_heading("谈判要点", level=2)
        nego = user_context.get("negotiation_points") or {}
        must = nego.get("must_win") or ["实现债权费用条款", "争议解决与送达条款"]
        should = nego.get("should_win") or ["违约责任细化", "验收标准完善"]
        p = doc.add_paragraph()
        self._label_run(p, "必守：", CLR_HIGH)
        p.add_run("；".join(must))
        p = doc.add_paragraph()
        self._label_run(p, "争取：", CLR_MID)
        p.add_run("；".join(should))
        doc.add_page_break()

        # ------------------------------------------------------------------
        # （一）风险总览仪表盘
        # ------------------------------------------------------------------
        doc.add_heading("（一）风险总览", level=1)

        table = doc.add_table(rows=4, cols=3)
        table.style = "Table Grid"
        heads = ["风险等级", "数量", "是否需立即处理"]
        for j, h in enumerate(heads):
            cell = table.rows[0].cells[j]
            self._shade(cell, SHADE_HEAD)
            r = cell.paragraphs[0].add_run(h)
            r.font.bold = True
        rows = [
            ("高", summary.get("致命风险", 0), "是" if summary.get("致命风险", 0) > 0 else "—", CLR_HIGH),
            ("中", summary.get("重要风险", 0), "视情况" if summary.get("重要风险", 0) else "—", CLR_MID),
            ("低", summary.get("一般风险", 0) + summary.get("轻微瑕疵", 0), "否", CLR_LOW),
        ]
        for i, (lv, count, action, color) in enumerate(rows, 1):
            cells = table.rows[i].cells
            self._label_run(cells[0].paragraphs[0], lv, color)
            cells[1].paragraphs[0].add_run(str(count))
            cells[2].paragraphs[0].add_run(action)

        doc.add_heading("八维度风险评分（1-5，5 为最高风险）", level=2)
        dim_scores = scoring.get("dimension_scores", {})
        for dim in radar_data:
            score = radar_data[dim]
            bar = "█" * int(score) + "░" * (5 - int(score))
            p = doc.add_paragraph()
            p.add_run(f"{dim}：{bar} ")
            label = dim_scores.get(dim, {}).get("label", "")
            self._label_run(p, f"{score}/5 {label}",
                            CLR_HIGH if score >= 4 else CLR_MID if score >= 3 else CLR_LOW)
        grade = scoring.get("risk_grade") or self._comprehensive_level(risk_report)
        p = doc.add_paragraph()
        self._label_run(p, f"综合风险等级：{grade}", CLR_HEADING)
        if scoring.get("deep_review_required"):
            p = doc.add_paragraph()
            self._label_run(p, "⚠ 存在评分 ≥4 的维度，建议律师深度审阅。", CLR_HIGH)

        if label_dist:
            doc.add_heading("风险类型分布", level=2)
            for label, count in sorted(label_dist.items(), key=lambda x: x[1], reverse=True):
                doc.add_paragraph(f"{label}：{count} 项", style="List Bullet")

        # ------------------------------------------------------------------
        # （二）合同基本信息
        # ------------------------------------------------------------------
        doc.add_heading("（二）合同基本信息", level=1)
        self._meta_card(doc, [
            ("合同类型", analysis_result.get("identified_type", "未知")),
            ("起草方", user_context.get("party", "未识别")),
            ("审核深度", user_context.get("review_depth", "标准审核")),
        ])

        # ------------------------------------------------------------------
        # （三）逐条审核意见（七列表格）
        # ------------------------------------------------------------------
        doc.add_heading("（三）逐条审核意见", level=1)
        cols = ["序号", "位置", "风险类型", "被审条款原文", "风险描述", "修改建议", "风险等级"]
        widths = [1.0, 1.8, 1.8, 3.6, 3.2, 3.2, 1.4]  # 内容列加宽
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        for j, (h, w) in enumerate(zip(cols, widths)):
            cell = table.rows[0].cells[j]
            cell.width = Cm(w)
            self._shade(cell, SHADE_HEAD)
            r = cell.paragraphs[0].add_run(h)
            r.font.bold = True

        idx = 1
        for level in ["致命风险", "重要风险", "一般风险", "轻微瑕疵"]:
            for risk in risks_by_level.get(level, []):
                cells = table.add_row().cells
                cells[0].paragraphs[0].add_run(str(idx))
                cells[1].paragraphs[0].add_run(risk.get("location", "—")[:20])
                cells[2].paragraphs[0].add_run(risk.get("risk_label", "—")[:12])
                cells[3].paragraphs[0].add_run(risk.get("original_text", "")[:60])
                cells[4].paragraphs[0].add_run(risk.get("description", "")[:80])
                cells[5].paragraphs[0].add_run(risk.get("suggestion", "")[:80])
                self._label_run(cells[6].paragraphs[0],
                                LEVEL_TEXT.get(level, "低"), LEVEL_COLOR.get(level, CLR_LOW))
                idx += 1

        # ------------------------------------------------------------------
        # （四）总体评价与签约利弊分析
        # ------------------------------------------------------------------
        doc.add_heading("（四）总体评价与签约利弊分析", level=1)

        doc.add_heading("有利因素", level=2)
        for item in user_context.get("pros") or ["合同类型明确，交易结构清晰", "核心条款基本覆盖交易关键环节"]:
            doc.add_paragraph(item, style="List Bullet")

        doc.add_heading("不利因素与重大风险提示", level=2)
        doc.add_paragraph(
            f"共发现 {risk_report.get('total_risks', 0)} 个风险点，"
            f"其中高风险 {summary.get('致命风险', 0)} 个、中风险 {summary.get('重要风险', 0)} 个。")
        for risk in risks_by_level.get("致命风险", []):
            p = doc.add_paragraph(style="List Bullet")
            self._label_run(p, risk.get("description", "")[:60] + "：", CLR_HIGH)
            p.add_run(risk.get("impact", "")[:60])

        doc.add_heading("谈判建议（三层优先级）", level=2)
        p = doc.add_paragraph()
        self._label_run(p, "第一层 必须获得（deal-breaker）：", CLR_HIGH)
        p.add_run("；".join(must))
        p = doc.add_paragraph()
        self._label_run(p, "第二层 应当获得：", CLR_MID)
        p.add_run("；".join(should))
        p = doc.add_paragraph()
        self._label_run(p, "第三层 可让步候选：", CLR_LOW)
        p.add_run("；".join(nego.get("nice_to_have") or ["格式统一", "文字优化"]))
        fallback = nego.get("fallback")
        if fallback:
            doc.add_paragraph(f"备用立场：{fallback}")

        # "清单之外的一个问题"强制追问（B6）
        doc.add_heading("清单之外的一个问题（强制追问）", level=2)
        beyond = user_context.get("beyond_checklist")
        if beyond:
            doc.add_paragraph(beyond)
        else:
            p = doc.add_paragraph()
            self._label_run(p, "【待律师补充】", CLR_LABEL)
            doc.add_paragraph(
                "从以下四个角度追问一个审核清单未覆盖、但可能影响客户决策的问题："
                "① 商业模式变化（客户业务调整后本合同是否仍适配）；"
                "② 对方履约能力恶化（资信/涉诉/股权变动）；"
                "③ 行业惯例（本类交易惯常做法与本文本的偏离）；"
                "④ 连锁效应（本合同对客户其他合同/融资/担保安排的影响）。")

        doc.add_heading("签署后注意事项", level=2)
        for item in ["妥善保存合同签署版本及往来沟通记录",
                     "注意合同约定的通知期限和异议期限",
                     "关注对方重大事项变化对合同履行的影响"]:
            doc.add_paragraph(item, style="List Bullet")

        # ------------------------------------------------------------------
        # （五）法律依据清单
        # ------------------------------------------------------------------
        doc.add_heading("（五）法律依据清单", level=1)
        legal_basis_items = []
        for level in ["致命风险", "重要风险", "一般风险", "轻微瑕疵"]:
            for risk in risks_by_level.get(level, []):
                basis = risk.get("legal_basis", "").strip()
                if basis and basis not in legal_basis_items:
                    legal_basis_items.append(basis)
        if legal_basis_items:
            for item in legal_basis_items:
                doc.add_paragraph(item, style="List Number")
        else:
            doc.add_paragraph("见逐条审核意见表中的依据列。")

        doc.add_paragraph()
        doc.add_paragraph(f"审核律师：{author}")
        doc.add_paragraph(f"审核日期：{today}")
        p = doc.add_paragraph()
        self._label_run(
            p,
            "声明：本意见书制作过程中使用了人工智能工具辅助起草，全部内容均经承办律师逐项审核、"
            "修改与确认，律师对文件内容承担全部专业责任。律师分析利弊，最终决策权在客户。",
            CLR_LABEL, bold=False)

        doc.save(str(filepath))
        print(f"法律意见书已生成: {filepath}")
        return str(filepath)

    # =======================================================================
    # 2. 法律分析（内部参考）
    # =======================================================================

    def generate_legal_analysis_docx(self, contract_name: str, analysis_result: Dict,
                                      risk_report: Dict) -> str:
        """生成法律分析 .docx（内部参考文件）"""
        filepath = self.output_dir / f"{contract_name}-法律分析.docx"
        risks_by_level = risk_report.get("risks_by_level", {})

        doc = self._base_doc()
        doc.add_heading(f"{contract_name} — 法律分析（内部参考）", level=0)
        self._meta_card(doc, [
            ("生成日期", datetime.now().strftime("%Y年%m月%d日")),
            ("合同类型", analysis_result.get("identified_type", "未知")),
            ("文件性质", "内部参考，不对客交付"),
        ])

        doc.add_heading("修订点对应法律依据", level=1)
        for level in ["致命风险", "重要风险", "一般风险", "轻微瑕疵"]:
            for risk in risks_by_level.get(level, []):
                doc.add_heading(f"[{level}] {risk.get('description', '')[:50]}", level=2)
                doc.add_paragraph(f"法律依据：{risk.get('legal_basis', '待补充')}")
                doc.add_paragraph(f"修改建议：{risk.get('suggestion', '')}")
                doc.add_paragraph(f"影响分析：{risk.get('impact', '')}")

        doc.add_heading("检索验证记录", level=1)
        p = doc.add_paragraph()
        self._label_run(
            p,
            "【待补充】各法律依据的检索验证状态（优先源命中数、充分性结论、失效复核结果），"
            "来源标注遵循分层检索协议。",
            CLR_LABEL, bold=False)

        doc.save(str(filepath))
        print(f"法律分析已生成: {filepath}")
        return str(filepath)

    # =======================================================================
    # 3. 批注版合同（双引擎：docx skill 增强 / ooxml_lite 内置兜底）
    # =======================================================================

    def generate_annotated_contract_docx(self, contract_name: str, original_docx_path: str,
                                          risk_report: Dict, author: str = "陈石律师【海泰所】",
                                          initials: str = "CS") -> Optional[str]:
        """生成批注版合同 .docx"""
        return self.generate_tracked_changes_docx(
            contract_name, original_docx_path, risk_report, author, initials
        )

    def generate_tracked_changes_docx(
        self,
        contract_name: str,
        original_docx_path: str,
        risk_report: Dict,
        author: str = "陈石律师【海泰所】",
        initials: str = "CS",
    ) -> Optional[str]:
        """生成带 Word 修订模式（Track Changes）和批注的 .docx 文件。

        引擎分派：环境中存在 docx skill 时使用其 Document library，
        否则使用内置 ooxml_lite 引擎（零依赖兜底）。
        """
        try:
            from docx_generator import DocxTrackChangesGenerator

            generator = DocxTrackChangesGenerator(
                original_docx_path=original_docx_path,
                risk_report=risk_report,
                output_dir=str(self.output_dir),
                author=author,
                initials=initials,
            )
            result = generator.generate(contract_name)
            if result:
                print(f"✅ 批注版合同已生成: {result}")
            else:
                print("⚠️ 批注版生成失败，已跳过")
            return result
        except Exception as e:
            print(f"⚠️ 批注版生成失败: {e}")
            return None

    # =======================================================================
    # 工具
    # =======================================================================

    @staticmethod
    def _comprehensive_level(risk_report: Dict) -> str:
        summary = risk_report.get("summary", {})
        fatal = summary.get("致命风险", 0)
        important = summary.get("重要风险", 0)
        if fatal > 0:
            return "极高风险"
        if important > 3:
            return "高风险"
        if important > 0:
            return "中风险"
        return "低风险"
